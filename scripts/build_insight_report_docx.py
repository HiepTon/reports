#!/usr/bin/env python3
"""Build MD OT Security v4.0.0 insight DOCX from merged Jira export + optional supplemental issues."""

from __future__ import annotations

import json
import os
import re
import statistics
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

os.environ.setdefault("MPLCONFIGDIR", str(Path(__file__).resolve().parents[1] / ".mpl"))

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / "data"
DEFAULT_ISSUES = DATA / "jira_mdots_4_0_0_issues.json"
SUPP = DATA / "supplemental_issues.json"
OUT = ROOT / "output" / "MD_OT_Security_v4.0.0_Insight_Report.docx"
CHART_DIR = ROOT / "output" / "_charts"

CONFLUENCE_REF = (
    "https://opswat.atlassian.net/wiki/spaces/NAC3/pages/5180621046/Test+Report+v4.0.0"
)

ROOT_BUCKET = ("ROOT", "NO_PARENT")

STOPWORDS = frozenset(
    """
    the a an and or to of in is was are be been being have has had having do does did doing
    will would could should may might must can for from with without into onto upon at by
    as if then than that this these those it its we our you your they their them not no
    when where what which who how why all any each every both few more most other such same
    so than too very just also only even ever still just about after before again here there
    actual expected steps observer step click navigate try use using used using user users
    http https com org net jira atlassian opswat blob null width height type file localid
    """.split()
)

# (label, keywords) — first match wins on ties by list order after scoring
BUG_AREA_RULES: list[tuple[str, tuple[str, ...]]] = [
    (
        "Import / CSV / SCD / assets",
        (
            "import", "csv", "scd", "batch", "asset", "field", "validation", "origin",
            "nic", "purdue", "manufacturer", "subtype", "type", "device when importing",
        ),
    ),
    (
        "Connection / Network / Policy",
        (
            "connection", "network map", "allowlist", "policy", "handshake", "sensor",
            "site bundle", "aio", "sync", "disappear", "inactive", "connection list",
        ),
    ),
    ("Dashboard / UI / filters", ("dashboard", "filter", "table", "loading", "display", "ui", "role", "checkbox")),
    ("Reporting / PDF / support package", ("pdf", "report", "support package", "export", "vulnerability")),
    ("Device / detection / classification", ("device", "detection", "classification", "mac", "cdp", "cve", "firmware")),
    ("Patching / Suricata / services", ("patch", "suricata", "toggle", "service", "worker", "queue")),
    ("DevOps / install / upgrade / CI", ("install", "upgrade", "nginx", "debian", "build", "pipeline", "runner", "hardware", "disk")),
    ("Alerts / notifications", ("alert", "acknowledge", "notification")),
    ("AI / ML services", ("ai", "library", "model", "inference")),
    ("Security / auth", ("auth", "permission", "token", "ssl", "tls")),
]

# Text-inferred component (engineering slice), aligned with summary tags + keywords
COMPONENT_RULES: list[tuple[str, tuple[str, ...]]] = [
    ("Frontend (text)", ("dashboard", "filter", "table", "ui", "page", "render", "component", "checkbox", "tooltip")),
    ("Backend (text)", ("api", "database", "sync", "device", "sensor", "connection", "policy", "server", "entity")),
    ("DevOps (text)", ("install", "upgrade", "nginx", "build", "pipeline", "runner", "debian", "hardware", "ci")),
    ("AI (text)", ("ai", "library", "model", "inference", "gpu")),
]


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, headers: list[str], rows: list[list[Any]], header_fill: str = "D9E2F3") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        set_cell_shading(hdr_cells[i], header_fill)
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = "" if val is None else str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph("")


def display_user(u: Any) -> str:
    if not u:
        return "Unassigned"
    if isinstance(u, dict):
        return u.get("displayName") or u.get("emailAddress") or "Unknown"
    return str(u)


def num(v: Any) -> float:
    if v is None:
        return 0.0
    if isinstance(v, (int, float)):
        return float(v)
    return 0.0


def option_value(v: Any) -> Optional[str]:
    if v is None:
        return None
    if isinstance(v, dict) and "value" in v:
        return str(v["value"])
    if isinstance(v, list) and v:
        parts = []
        for x in v:
            if isinstance(x, dict) and x.get("value") is not None:
                parts.append(str(x["value"]))
        return ", ".join(parts) if parts else None
    return str(v)


def parse_dt(s: Optional[str]) -> Optional[datetime]:
    if not s:
        return None
    try:
        return datetime.strptime(s, "%Y-%m-%dT%H:%M:%S.%f%z")
    except ValueError:
        pass
    try:
        return datetime.strptime(s, "%Y-%m-%dT%H:%M:%S%z")
    except ValueError:
        pass
    return None


def days_between(a: Optional[datetime], b: Optional[datetime]) -> Optional[float]:
    if not a or not b:
        return None
    return abs((b - a).total_seconds() / 86400.0)


def summary_lane(summary: str) -> str:
    s = summary or ""
    if re.match(r"^\s*\[FE\]", s, re.I):
        return "Frontend"
    if re.match(r"^\s*\[BE\]", s, re.I):
        return "Backend"
    if re.match(r"^\s*\[DevOps\]", s, re.I):
        return "DevOps"
    return "Other"


def strip_jira_markup(raw: Optional[str]) -> str:
    if not raw:
        return ""
    t = str(raw)
    t = re.sub(r"<[^>]+>", " ", t)
    t = re.sub(r"\[([^\]|]+)\|([^\]]+)\]", r"\2", t)
    t = re.sub(r"\s+", " ", t).strip()
    return t


def issue_text_blob(issue: dict[str, Any]) -> str:
    f = issue["fields"]
    parts = [f.get("summary") or ""]
    parts.append(strip_jira_markup(f.get("description")))
    ac = f.get("customfield_16728")
    if ac:
        parts.append(strip_jira_markup(ac))
    return " ".join(parts).lower()


def score_labels(text: str, rules: list[tuple[str, tuple[str, ...]]]) -> tuple[str, int]:
    best_label = "Uncategorized / general"
    best_score = 0
    for label, kws in rules:
        s = sum(text.count(kw) for kw in kws)
        if s > best_score:
            best_score = s
            best_label = label
    if best_score == 0:
        return best_label, 0
    return best_label, best_score


def infer_bug_area(issue: dict[str, Any]) -> str:
    return score_labels(issue_text_blob(issue), BUG_AREA_RULES)[0]


def infer_text_component(issue: dict[str, Any]) -> str:
    summary = issue["fields"].get("summary") or ""
    sl = summary_lane(summary)
    if sl == "Frontend":
        return "Frontend (from [FE] tag)"
    if sl == "Backend":
        return "Backend (from [BE] tag)"
    if sl == "DevOps":
        return "DevOps (from [DevOps] tag)"
    blob = issue_text_blob(issue)
    label, sc = score_labels(blob, COMPONENT_RULES)
    if sc == 0:
        return "General / mixed (text)"
    return label


def first_sentence(text: str, max_len: int = 320) -> str:
    if not text.strip():
        return ""
    t = text.strip()
    for sep in (". ", "? ", "! ", "\n"):
        if sep in t[:600]:
            t = t.split(sep)[0] + sep.strip()
            break
    if len(t) > max_len:
        t = t[: max_len - 1].rsplit(" ", 1)[0] + "…"
    return t


def extract_key_stories(issues: list[dict[str, Any]], limit: int = 14) -> list[list[str]]:
    """Prioritize Stories (and Epics with narrative) by story points and description richness."""
    rows: list[tuple[float, int, dict[str, Any]]] = []
    for i in issues:
        it = i["fields"]["issuetype"]["name"]
        if it not in ("Story", "Epic"):
            continue
        pts = num(i["fields"].get("customfield_16486"))
        desc = strip_jira_markup(i["fields"].get("description"))
        richness = len(desc) + (50 if it == "Epic" else 0)
        rows.append((pts, richness, i))
    rows.sort(key=lambda x: (-x[0], -x[1], x[2]["key"]))
    out: list[list[str]] = []
    for pts, _rich, i in rows[:limit]:
        plain = strip_jira_markup(i["fields"].get("description"))
        narrative = first_sentence(plain) if plain else "(No description — see summary.)"
        if len(narrative) < 20:
            narrative = first_sentence(i["fields"].get("summary") or "", 200)
        out.append(
            [
                i["key"],
                i["fields"]["issuetype"]["name"],
                (i["fields"].get("summary") or "")[:100],
                str(int(pts)) if pts == int(pts) else str(pts) if pts else "",
                narrative[:400],
            ]
        )
    return out


def word_freq(texts: list[str], top_n: int = 12) -> list[tuple[str, int]]:
    blob = " ".join(texts).lower()
    words = re.findall(r"[a-z][a-z0-9_-]{3,}", blob)
    c: Counter[str] = Counter(w for w in words if w not in STOPWORDS)
    return c.most_common(top_n)


def generate_lane_insights(lane: str, lane_issues: list[dict[str, Any]]) -> list[str]:
    if not lane_issues:
        return ["No bugs or enhancements in this lane for the selected fix version."]
    texts = [issue_text_blob(i) for i in lane_issues]
    pri = Counter((i["fields"].get("priority") or {}).get("name") or "—" for i in lane_issues)
    st = Counter((i["fields"].get("status") or {}).get("name") or "—" for i in lane_issues)
    bullets: list[str] = []
    bullets.append(
        f"{len(lane_issues)} bug-like tickets; dominant priority: {pri.most_common(1)[0][0]} "
        f"({pri.most_common(1)[0][1]} tickets); most common status: {st.most_common(1)[0][0]}."
    )
    areas = Counter(infer_bug_area(i) for i in lane_issues)
    bullets.append(
        "Description-inferred hotspots: "
        + ", ".join(f"{a} ({c})" for a, c in areas.most_common(4))
        + "."
    )
    for w, c in word_freq(texts, 8)[:5]:
        if c < 2:
            continue
        bullets.append(f"Recurring language around “{w}” ({c} mentions) — review tests and docs in that subsystem.")
    sev = pri.get("Blocker", 0) + pri.get("Critical", 0) + pri.get("Major", 0)
    if sev:
        bullets.append(f"{sev} ticket(s) are Major or above; ensure targeted regression and sign-off evidence before release.")
    return bullets


def save_horizontal_bar(
    path: Path,
    title: str,
    labels: list[str],
    values: list[float],
    xlabel: str,
) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    h = max(2.8, 0.38 * len(labels) + 1.2)
    fig, ax = plt.subplots(figsize=(8.5, h))
    y = list(range(len(labels)))
    ax.barh(y, values, color="#4472C4")
    ax.set_yticks(y)
    ax.set_yticklabels(labels, fontsize=9)
    ax.set_xlabel(xlabel)
    ax.set_title(title, fontsize=11, pad=10)
    ax.invert_yaxis()
    fig.tight_layout()
    fig.savefig(path, dpi=120, bbox_inches="tight")
    plt.close(fig)


def add_chart_paragraph(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(6.4))
    cap = doc.add_paragraph(caption)
    for r in cap.runs:
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def load_issues() -> list[dict[str, Any]]:
    issues = json.loads(DEFAULT_ISSUES.read_text())["issues"]
    if SUPP.exists():
        extra = json.loads(SUPP.read_text())
        if isinstance(extra, list):
            keys = {i["key"] for i in issues}
            for i in extra:
                if i["key"] not in keys:
                    issues.append(i)
                    keys.add(i["key"])
    return issues


def resolve_feature_bucket(
    issue: dict[str, Any], by_key: dict[str, dict[str, Any]]
) -> tuple[str, str, str, Optional[str], Optional[str]]:
    fields = issue["fields"]
    visited: set[str] = set()
    cur_key = issue["key"]

    while True:
        inf = fields.get("issuetype", {}).get("name")
        if inf == "Epic":
            drv = option_value(fields.get("customfield_17048"))
            pri = (fields.get("priority") or {}).get("name")
            return "EPIC", cur_key, fields.get("summary") or cur_key, drv, pri

        p = fields.get("parent")
        if not p:
            break
        pk = p["key"]
        if pk in visited:
            break
        visited.add(pk)

        pit = (p.get("fields") or {}).get("issuetype", {}).get("name")
        if pit == "Epic":
            pf = p["fields"]
            epic = by_key.get(pk)
            if epic:
                ef = epic["fields"]
                drv = option_value(ef.get("customfield_17048"))
                pri = (ef.get("priority") or {}).get("name")
                title = ef.get("summary") or pk
            else:
                drv = option_value(pf.get("customfield_17048")) if isinstance(pf, dict) else None
                pri = (pf.get("priority") or {}).get("name") if isinstance(pf, dict) else None
                title = pf.get("summary") if isinstance(pf, dict) else pk
            return "EPIC", pk, title or pk, drv, pri

        nxt = by_key.get(pk)
        if nxt:
            cur_key = pk
            fields = nxt["fields"]
            continue

        pf = p.get("fields") or {}
        title = pf.get("summary") or pk
        pri = (pf.get("priority") or {}).get("name")
        return "STORY", pk, f"{title} ({pk})", None, pri

    if fields.get("issuetype", {}).get("name") == "Epic":
        drv = option_value(fields.get("customfield_17048"))
        pri = (fields.get("priority") or {}).get("name")
        return "EPIC", issue["key"], fields.get("summary") or issue["key"], drv, pri

    return "ROOT", ROOT_BUCKET[1], "Root-level items (no parent in Jira)", None, None


def primary_role(lanes: Counter, types: Counter) -> str:
    if lanes:
        lane, n = lanes.most_common(1)[0]
        if n > 0 and lane != "Other":
            return lane
    if not types:
        return "Contributor"
    t, _ = types.most_common(1)[0]
    if t == "Bug":
        return "Engineering / QA"
    if t in ("Story", "Epic", "Task"):
        return "Engineering"
    if t == "Subtask":
        return "Engineering"
    if t == "Enhancement":
        return "Engineering / Product"
    return str(t)


def main() -> None:
    issues = load_issues()
    by_key: dict[str, dict[str, Any]] = {i["key"]: i for i in issues}

    bucket_sums: dict[tuple[str, str], dict[str, float]] = defaultdict(
        lambda: {"dev": 0.0, "qa": 0.0, "rev": 0.0, "pts": 0.0}
    )
    bucket_meta: dict[tuple[str, str], tuple[str, Optional[str], Optional[str]]] = {}

    for i in issues:
        it = i["fields"]["issuetype"]["name"]
        if it == "Epic":
            k = ("EPIC", i["key"])
            bucket_meta[k] = (
                i["fields"].get("summary") or i["key"],
                option_value(i["fields"].get("customfield_17048")),
                (i["fields"].get("priority") or {}).get("name"),
            )
            continue

        kind, bid, title, drv, pri = resolve_feature_bucket(i, by_key)
        bkey = (kind, bid)
        bucket_sums[bkey]["dev"] += num(i["fields"].get("customfield_21867"))
        bucket_sums[bkey]["qa"] += num(i["fields"].get("customfield_21869"))
        bucket_sums[bkey]["rev"] += num(i["fields"].get("customfield_21868"))
        bucket_sums[bkey]["pts"] += num(i["fields"].get("customfield_16486"))

        if bkey not in bucket_meta:
            bucket_meta[bkey] = (title, drv, pri)
        else:
            old_title, old_drv, old_pri = bucket_meta[bkey]
            if not old_drv and drv:
                bucket_meta[bkey] = (old_title, drv, old_pri or pri)

    for i in issues:
        if i["fields"]["issuetype"]["name"] == "Epic":
            ek = ("EPIC", i["key"])
            bucket_sums.setdefault(ek, {"dev": 0.0, "qa": 0.0, "rev": 0.0, "pts": 0.0})

    feature_rows: list[list[Any]] = []
    for bkey in sorted(bucket_sums.keys(), key=lambda x: (x[0] != "EPIC", x[0] != "STORY", x[1])):
        sums = bucket_sums[bkey]
        total = sums["dev"] + sums["qa"] + sums["rev"]
        title, drv, pri = bucket_meta.get(bkey, (bkey[1], None, None))
        feature_rows.append(
            [
                title,
                round(sums["pts"], 2) if sums["pts"] else "",
                round(sums["dev"], 2) if sums["dev"] else "",
                round(sums["qa"], 2) if sums["qa"] else "",
                round(sums["rev"], 2) if sums["rev"] else "",
                round(total, 2) if total else "",
                drv or "",
                pri or "",
            ]
        )

    member_types: dict[str, Counter[str]] = defaultdict(Counter)
    member_lanes: dict[str, Counter[str]] = defaultdict(Counter)
    member_story_pts: dict[str, float] = defaultdict(float)
    member_bugs: dict[str, int] = defaultdict(int)
    for i in issues:
        assignee = display_user(i["fields"].get("assignee"))
        it = i["fields"]["issuetype"]["name"]
        member_types[assignee][it] += 1
        member_lanes[assignee][summary_lane(i["fields"].get("summary") or "")] += 1
        member_story_pts[assignee] += num(i["fields"].get("customfield_16486"))
        if it in ("Bug", "Enhancement"):
            member_bugs[assignee] += 1

    def pct(part: float, whole: float) -> str:
        if whole <= 0:
            return ""
        return f"{100.0 * part / whole:.1f}%"

    issue_assigned_total = float(sum(sum(member_types[m].values()) for m in member_types))
    team_story_pts = float(sum(member_story_pts.values()))

    def primary_focus_sentence(_m: str, tc: Counter[str], lanes: Counter[str]) -> str:
        lane = primary_role(lanes, tc)
        parts: list[str] = []
        if lane == "Frontend":
            parts.append("Frontend")
        elif lane == "Backend":
            parts.append("Backend")
        elif lane == "DevOps":
            parts.append("DevOps")
        else:
            parts.append(lane)
        if tc.get("Bug", 0) + tc.get("Enhancement", 0) >= max(2, tc.get("Story", 0)):
            parts.append("bugs & fixes")
        elif tc.get("Story", 0) or tc.get("Task", 0):
            parts.append("features & tasks")
        if tc.get("Subtask", 0) >= 3:
            parts.append("subtasks")
        return ": ".join(parts[:2]) if len(parts) > 1 else (parts[0] if parts else "—")

    member_rows = []
    for m in member_types.keys():
        tc = member_types[m]
        tasks = tc.get("Task", 0)
        subs = tc.get("Subtask", 0)
        stories = tc.get("Story", 0)
        bugs = tc.get("Bug", 0) + tc.get("Enhancement", 0)
        total = sum(tc.values())
        role = primary_role(member_lanes[m], tc)
        sp = member_story_pts[m]
        focus = primary_focus_sentence(m, tc, member_lanes[m])
        member_rows.append(
            [
                m,
                round(sp, 2) if sp else "",
                pct(sp, team_story_pts),
                role,
                focus,
                tasks,
                subs,
                stories,
                bugs,
                total,
                pct(float(total), issue_assigned_total),
            ]
        )

    member_rows.sort(
        key=lambda r: (
            -float(r[1]) if r[1] != "" else 0.0,
            -int(r[9]),
            str(r[0]),
        )
    )

    bug_like = [i for i in issues if i["fields"]["issuetype"]["name"] in ("Bug", "Enhancement")]

    cat_counter_jira: Counter[str] = Counter()
    for i in bug_like:
        raw = i["fields"].get("customfield_17312")
        if isinstance(raw, list) and raw:
            for o in raw:
                if isinstance(o, dict) and o.get("value"):
                    cat_counter_jira[str(o["value"])] += 1
        elif isinstance(raw, dict) and raw.get("value"):
            cat_counter_jira[str(raw["value"])] += 1
        else:
            cat_counter_jira["(Uncategorized in Jira)"] += 1

    cat_counter_text: Counter[str] = Counter()
    for i in bug_like:
        cat_counter_text[infer_bug_area(i)] += 1

    comp_counter_jira: Counter[str] = Counter()
    for i in issues:
        comps = i["fields"].get("components") or []
        if not comps:
            comp_counter_jira["(No Jira component)"] += 1
        else:
            for c in comps:
                nm = c.get("name") if isinstance(c, dict) else str(c)
                comp_counter_jira[nm or "(None)"] += 1

    comp_counter_text: Counter[str] = Counter()
    for i in issues:
        comp_counter_text[infer_text_component(i)] += 1

    major_bugs = [
        i
        for i in bug_like
        if (i["fields"].get("priority") or {}).get("name") in ("Major", "Critical", "Blocker")
    ]
    major_rows = [
        [
            i["key"],
            (i["fields"].get("summary") or "")[:120],
            (i["fields"].get("status") or {}).get("name"),
            (i["fields"].get("priority") or {}).get("name"),
            display_user(i["fields"].get("assignee")),
        ]
        for i in sorted(major_bugs, key=lambda x: x["key"])
    ]

    res_days: list[float] = []
    for i in bug_like:
        rd = i["fields"].get("resolutiondate")
        cr = i["fields"].get("created")
        d0, d1 = parse_dt(cr), parse_dt(rd)
        dd = days_between(d0, d1)
        if dd is not None:
            res_days.append(dd)

    def fmt_stats(vals: list[float]) -> str:
        if not vals:
            return "No resolved items with resolution timestamps in this export."
        vals_sorted = sorted(vals)
        med = statistics.median(vals_sorted)
        avg = statistics.mean(vals_sorted)
        p90 = vals_sorted[max(0, int(0.9 * (len(vals_sorted) - 1)))]
        return f"Count: {len(vals)} | Mean: {avg:.1f} d | Median: {med:.1f} d | P90: {p90:.1f} d"

    bug_assignee: Counter[str] = Counter()
    for i in bug_like:
        bug_assignee[display_user(i["fields"].get("assignee"))] += 1
    bug_assignee_rows = [[a, c] for a, c in bug_assignee.most_common()]

    def bug_detail_rows(lane: str) -> list[list[str]]:
        rows: list[list[str]] = []
        for i in bug_like:
            s = i["fields"].get("summary") or ""
            if summary_lane(s) != lane:
                continue
            rows.append(
                [
                    i["key"],
                    s[:140],
                    (i["fields"].get("status") or {}).get("name") or "",
                    (i["fields"].get("priority") or {}).get("name") or "",
                    display_user(i["fields"].get("assignee")),
                ]
            )
        return sorted(rows, key=lambda x: x[0])

    def lane_issue_list(lane: str) -> list[dict[str, Any]]:
        return [i for i in bug_like if summary_lane(i["fields"].get("summary") or "") == lane]

    key_story_rows = extract_key_stories(issues)

    CHART_DIR.mkdir(parents=True, exist_ok=True)
    chart_sp = CHART_DIR / "member_story_points.png"
    chart_bugs = CHART_DIR / "member_bugs.png"
    top_members = [r[0] for r in member_rows[:16]]
    sp_vals = [float(r[1]) if r[1] != "" else 0.0 for r in member_rows[:16]]
    bug_vals = [float(member_bugs.get(r[0], 0)) for r in member_rows[:16]]
    save_horizontal_bar(
        chart_sp,
        "Story points by assignee (top contributors)",
        top_members,
        sp_vals,
        "Sum of Jira Story point estimate",
    )
    save_horizontal_bar(
        chart_bugs,
        "Bugs + enhancements by assignee (same order as table)",
        top_members,
        bug_vals,
        "Ticket count",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h = doc.add_heading("MD OT Security v4.0.0 — Insight Report", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph()
    p.add_run("Source: ").bold = True
    p.add_run(CONFLUENCE_REF)
    doc.add_paragraph(
        "Generated from Jira project OTSECURITY, fix version “MD OT Security 4.0.0”. "
        "Summaries matching “[Cancel]” were excluded where applicable. "
        "Several sections combine Jira fields with description-based inference (keyword scoring over summary + description + acceptance criteria when present)."
    )

    doc.add_heading("1. Key stories (from ticket content)", level=1)
    doc.add_paragraph(
        "Stories and epics ranked by story points, then by description length. "
        "The “Narrative (from description)” column is the first sentence (or lead fragment) of plain-text description."
    )
    add_table(
        doc,
        ["Key", "Type", "Summary", "Story pts", "Narrative (from description)"],
        key_story_rows,
    )

    doc.add_heading("2. Features (Epic / story scope)", level=1)
    add_table(
        doc,
        [
            "Feature / scope",
            "Story pts (Jira Σ)",
            "Dev estimate (Σ)",
            "QA estimate (Σ)",
            "Review estimate (Σ)",
            "Total est. (Σ)",
            "Driver",
            "Priority",
        ],
        feature_rows,
    )

    doc.add_heading("3. Member effort and contributions", level=1)
    doc.add_paragraph(
        "Story points sum the Jira “Story point estimate” on assigned issues. "
        "Percentages use team totals in this export. Role / primary focus use [FE]/[BE]/[DevOps] and issue-type mix."
    )
    add_chart_paragraph(
        doc,
        chart_sp,
        "Figure 1 — Story point concentration helps spot single points of failure in planning.",
    )
    add_chart_paragraph(
        doc,
        chart_bugs,
        "Figure 2 — Bug + enhancement counts by assignee (same top slice as the chart above).",
    )
    add_table(
        doc,
        [
            "Member",
            "Story points (Σ)",
            "% of team SP",
            "Role",
            "Primary focus",
            "Tasks",
            "Subtasks",
            "Stories",
            "Bugs + Enh.",
            "Total issues",
            "% of issues",
        ],
        member_rows,
    )

    doc.add_heading("4. Bug breakdown by area", level=1)
    doc.add_paragraph(
        "Two views: Jira Bug Category (when set) vs. description-driven buckets from keyword scoring over summary + description."
    )
    doc.add_paragraph("4a — Jira Bug Category", style="Heading 3")
    add_table(doc, ["Bug category (Jira)", "Count"], [[k, v] for k, v in cat_counter_jira.most_common()])
    doc.add_paragraph("4b — Inferred from ticket text", style="Heading 3")
    add_table(doc, ["Inferred area (text)", "Count"], [[k, v] for k, v in cat_counter_text.most_common()])

    doc.add_heading("5. Issues by component", level=1)
    doc.add_paragraph(
        "Jira component field (tickets may list multiple components). "
        "Text-inferred slice uses [FE]/[BE]/[DevOps] tags when present, else keyword rules on description."
    )
    doc.add_paragraph("5a — Jira components", style="Heading 3")
    add_table(doc, ["Component (Jira)", "Row count"], [[k, v] for k, v in comp_counter_jira.most_common()])
    doc.add_paragraph("5b — Text-inferred component", style="Heading 3")
    add_table(doc, ["Inferred component (text)", "Issue count"], [[k, v] for k, v in comp_counter_text.most_common()])

    doc.add_heading("6. Major bugs (Major / Critical / Blocker)", level=1)
    add_table(doc, ["Key", "Summary", "Status", "Priority", "Assignee"], major_rows)

    doc.add_heading("7. Resolution time analysis (Bugs + Enhancements)", level=1)
    doc.add_paragraph(fmt_stats(res_days))

    doc.add_heading("8. Workload distribution — bugs + enhancements by assignee", level=1)
    add_table(doc, ["Assignee", "Count"], bug_assignee_rows)

    for idx, (lane, title) in enumerate(
        [
            ("Frontend", "9. Frontend bug detail ([FE])"),
            ("Backend", "10. Backend bug detail ([BE])"),
            ("DevOps", "11. DevOps bug detail ([DevOps])"),
        ],
        start=0,
    ):
        doc.add_heading(title, level=1)
        for bullet in generate_lane_insights(lane, lane_issue_list(lane)):
            doc.add_paragraph(bullet, style="List Bullet")
        add_table(doc, ["Key", "Summary", "Status", "Priority", "Assignee"], bug_detail_rows(lane))

    doc.add_heading("12. Recommendations", level=1)
    recs = [
        "Link root-level stories and tasks to epics in Jira so release reporting rolls up drivers and priorities consistently.",
        "Keep Black Duck high-severity items on a versioned remediation plan (per release test report) instead of open-ended acknowledgement.",
        "Back UI and asset-scanning themes called out in sign-off with targeted regression automation for the next train.",
        "Enrich descriptions and Bug Category at filing time — text-inferred buckets are useful but should converge with structured fields.",
    ]
    for r in recs:
        doc.add_paragraph(r, style="List Bullet")

    doc.add_heading("13. Insights and key takeaways for the next release", level=1)
    n_sub = sum(1 for i in issues if i["fields"]["issuetype"]["name"] == "Subtask")
    insights = [
        f"Scope after exclusions: {len(issues)} issues; subtasks are a large share ({n_sub}), so epic/story health metrics should stay visible to leadership.",
        f"Bugs plus enhancements in this fix version: {len(bug_like)}; severity Major or above: {len(major_bugs)}.",
        "Compare mean versus median resolution time — a wide gap usually indicates a few long-tail defects that needed escalation or environment access.",
        "If bug workload is concentrated on a small set of assignees, plan load-balancing and earlier DevOps involvement for pipeline-impacting defects.",
        "Align security scanning sign-off with release criteria early to avoid late-cycle risk acceptance carrying into the next version.",
        "Compare Jira Bug Category vs. text-inferred areas: large gaps highlight tickets missing triage metadata.",
    ]
    for t in insights:
        doc.add_paragraph(t, style="List Bullet")

    doc.add_paragraph("")
    foot = doc.add_paragraph()
    run = foot.add_run(f"Generated automatically. Issue snapshot: {len(issues)} items.")
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    main()
